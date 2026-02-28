import docx
from docx.shared import Inches
import os
import glob
from pdf2image import convert_from_path, convert_from_bytes
from pdf2image.exceptions import (
    PDFInfoNotInstalledError,
    PDFPageCountError,
    PDFSyntaxError
)
# Get the current working directory
cwd = os.getcwd()
output_file = input("This script assumes that there is already a .docx created \n if this is the first time runing this script please first run the alpha_first_time_report_generator.py \n Enter the name of the already existing output file: ")
#output_file = "CleftPalete_WGS_GATK-SV"
# Print the current working directory
print("Current working directory: {0}".format(cwd))

#path_to_files = r'/Users/szaheri/Documents/SV/GATK-SV/github_ryan_shared/bwa_dragen_sample_level_comparison/results/bwa/perFamily_plots_tarball/bwa_samples_aligning_with_dragen_perFamily_plots/supporting_plots/sv_inheritance_plots/'
#folder = r'/Users/szaheri/Documents/SV/GATK-SV/github_ryan_shared/bwa_dragen_sample_level_comparison/results/bwa/perFamily_plots_tarball/bwa_samples_aligning_with_dragen_perFamily_plots/supporting_plots/sv_inheritance_plots/'

path_to_files = os.getcwd()+ '/'
folder = os.getcwd()+ '/'
head_tail = os.path.split(cwd)

print(head_tail)
# print tail part of the path
print(head_tail[1])

files = [folder + fn for fn in os.listdir(folder) if fn.endswith('.pdf')]
print(files)
# Convert them.
def pic(path, name):
	images = convert_from_path(path + name) #+ '.pdf'
	for i in range(len(images)):
		images[i].save(name[:-4] +str(i) +'.jpg', 'JPEG')


#pic(path_to_files,'all.pdf')


for filee in glob. glob("*.pdf"):
	#print(filee)
	pic(path_to_files,filee)



#mydoc = docx.Document()
mydoc = docx.Document(str(output_file) + ".docx")
mydoc.add_paragraph("This is a report of the " + str(head_tail[0]))
mydoc.save(str(output_file) + ".docx")


#mydoc.add_paragraph("This is the second paragraph of a MS Word file.")
#mydoc.save("my_written_file.docx")


files = [folder + fn for fn in os.listdir(folder) if fn.endswith('.pdf')]
print(files)
# Convert them.
#def pic(path, name):
#	images = convert_from_path(path + name) #+ '.pdf'
#	for i in range(len(images)):
#		images[i].save(name +str(i) +'.jpg', 'JPEG')


#pic(path_to_files,'all.pdf')

mydoc.add_heading(str(head_tail[1]), 0)
mydoc.save(str(output_file) + ".docx")

for filee in glob. glob("*.jpg" or "*.png"):
	print(filee)
	#pic(path_to_files,filee)
	mydoc.add_heading(str(filee[:-4]), 1)
	mydoc.save(str(output_file) + ".docx")
	mydoc.add_picture(filee, width=Inches(6.0))#, width=docx.shared.Inches(5), height=docx.shared.Inches(5))
	mydoc.save(str(output_file) + ".docx")


for filee in glob. glob("*.png"):
	print(filee)
	#pic(path_to_files,filee)
	mydoc.add_heading(str(filee[:-4]), 1)
	mydoc.save(str(output_file) + ".docx")
	mydoc.add_picture(filee, width=Inches(6.0))#, width=docx.shared.Inches(5), height=docx.shared.Inches(5))
	mydoc.save(str(output_file) + ".docx")
# You’re free to specify both width and height, but usually you wouldn’t want to. 
# If you specify only one, python-docx uses it to calculate the properly scaled value of the other. 
# This way the aspect ratio is preserved and your picture doesn’t look stretched.